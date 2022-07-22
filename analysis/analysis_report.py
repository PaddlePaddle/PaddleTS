# !/usr/bin/env python3
# -*- coding:utf-8 -*-

from io import BytesIO
import inspect
import json
import os
from typing import Any, Callable, List, Optional, Sequence, Tuple, Union, Dict

import pandas as pd
from docx import Document
from docx.shared import Inches

from paddlets import TimeSeries, TSDataset, analysis
from paddlets.analysis.base import Analyzer
from paddlets.logger import Logger, raise_if_not, raise_if, raise_log

logger = Logger(__name__)

# Default_analyzers
DEFAULT_ANALYZERS = ["summary", "max"]


class AnalysisReport(object):
    """
    AnalysisReport
    Aggregate the results of analyzers, show them in format of tables and charts in reports.
    Currently supprt docx and json Analysis report.
    
    Args:
        dataset[TSDataset]: TSDataset to be analyzed
        names[str|List[str]]: Analyzer names, set to DEFAULT_ANALYZERS by default
        params[Dict]: analyzers params
        columns[str|List[str]]: columns to be analyzed
    
    Examples:
        .. code-block:: python

            # example for names.
            names = ["max", "outlier"]

            # example for params:
            params = {
                "max":{
                "param1":1
                },
            "outlier":{
                "param1":1,
                "param2":2
                }}

    """

    def __init__(
            self,
            dataset: TSDataset,
            names: Union[str, List[str]] = None,
            params: Dict = None,
            columns: Optional[Union[str, List[str]]] = None
    ) -> None:

        if names == None:
            names = DEFAULT_ANALYZERS
        else:
            self._validate_analyzers_names(names)

        self._dataset = dataset
        self._names = names
        self._columns = columns
        self._analyzers = self._get_analyzers(names, params)

    def export_docx_report(self, path: str = ".", file_name: str = "analysis_report.docx") -> None:
        """
        Export a report in the docx format
        
        Args:
            path[str]: path to save the exported report, set to the current path by default
            file_name[str]: file name ,default set to "analysis_report.pdf"

        Returns:
            None

        """
        # Validate export path
        if not os.path.exists(path):
            raise_log(ValueError("export path do not exist, please check"))

        # Create new Document
        document = Document()

        self._report_formating(document)

        figure = self._dataset.plot().get_figure()
        if figure:
            memfile = BytesIO()
            figure.savefig(memfile)
            document.add_heading("Data View", level=2)
            document.add_picture(memfile, width=Inches(5.0))

        for analyzer in self._analyzers:
            properties = analyzer.get_properties()
            report_heading = properties.get("report_heading")
            report_description = properties.get("report_description")
            document.add_heading(report_heading, level=2)
            document.add_paragraph(report_description)
            
            # Display the analysis result
            document.add_paragraph("Analysis Results", style='ListBullet')
            analysis_result = analyzer(self._dataset, self._columns)
            if isinstance(analysis_result, pd.Series):
                analysis_result = pd.DataFrame(analysis_result)
            # Dataframe to table
            if isinstance(analysis_result, pd.DataFrame):
                indexes = analysis_result.index.to_list()
                t = document.add_table(analysis_result.shape[0]+1, analysis_result.shape[1]+1)
                # Add the header rows.
                t.cell(0,1).text = ""
                for j in range(analysis_result.shape[-1]):
                    t.cell(0,j+1).text = str(analysis_result.columns[j])
                    # Add the rest of the data frame
                for i in range(analysis_result.shape[0]):
                    t.cell(i+1,0).text = str(indexes[i])
                    for j in range(analysis_result.shape[-1]):
                        t.cell(i+1,j+1).text = str(analysis_result.values[i,j])
                t.style = "Table Grid"
            else: 
                document.add_paragraph(str(analysis_result))

            # Add figures
            figure = analyzer.plot()
            if figure:
                memfile = BytesIO()
                figure.savefig(memfile)
                document.add_paragraph("Charts", style='ListBullet')
                document.add_picture(memfile, width=Inches(5.0))

        document.add_page_break()
        path = path + "/" + file_name
        document.save(path)
        logger.info(f"save report succcess, save at {path}")
    
    def export_json_report(self, log: bool = True) -> Dict:
        """
        Export a report in the Json format
        
        Args:
            log[bool]: print log or not, default set to True

        Returns:
            Dict

        """
        json_report = {}

        for analyzer in self._analyzers:
            analyzer_report =  {}
            properties = analyzer.get_properties()
            report_heading = properties.get("report_heading")
            report_description = properties.get("report_description")

            analyzer_report["heading"] = report_heading
            analyzer_report["description"] = report_description

            analysis_res = analyzer(self._dataset, self._columns)

            if isinstance(analysis_res, pd.DataFrame) or isinstance(analysis_res, pd.Series):
                analyzer_report["analysis_results"] = analysis_res.to_json()
            else:
                analyzer_report["analysis_results"] = analysis_res

            analyzer_name = properties.get("name")
            json_report[analyzer_name] = analyzer_report 
                
        if log:
            logger.info(json_report)
        
        return json_report

    def _report_formating(self, document: Document) -> None:
        """
        Initialize analysis report 
        
        Args:
            document[Document]: document need to be initialized 

        Returns:
            None

        """
        document.add_heading(u' Data Analysis Report ', 0)
        # Add_aragraph
        document.add_paragraph(u'This report shows some analysis results in the form of tables and charts')
        document.add_paragraph(u'It is designed to give users a brief overview about the dataset')
        document.add_paragraph(u'Currently, the following analysis methods are supported, including:')
        # ListBullet
        document.add_paragraph(
            u'summary, max, fft, stft, cwt', style='ListBullet')

    def _get_analyzers(self, names: Union[str, List[str]], params: Dict = None) -> List[Analyzer]:
        """
        Get analyzer objects
        
        Args:
            names[str|List(str)]: analyzers name list, Not None
            params: the parameters of each analyzer

        Returns:
            List[Analyzer]

        Examples:
            .. code-block:: python
            example for params：
                   {
                    "max":{
                    "    param1":1
                        },
                    "outlier":{
                        "param1":1,
                        "param2":2
                    }}

        """
        analyzers = []
        analyzers_mapping = self._get_analyzers_mapping()
        for name, analyzer_obj in analyzers_mapping.items():
            if name not in names:
                continue
            param = params.get(name, None) if params is not None else None
            if param:
                analyzers.append(analyzer_obj(**param))
            else:
                analyzers.append(analyzer_obj())

        return analyzers

    def get_all_analyzers_names(self, log: bool = True) -> List[str]:
        """
        Get the names of analyzers
        This method can be called internally or externally, and the parameter log is set to False or True accordingly.
        
        Args:
            log(bool) : Whether to print the log, 
                        the default is True when used externally, and set to False when called internally

        Returns:
            List[str]
        """
        analyzers_mapping = self._get_analyzers_mapping()
        analyzers_names = []
        for key, value in analyzers_mapping.items():
            analyzers_names.append(key)
        if log:
            logger.info("current support analyzers:" + ','.join(analyzers_names))

        return analyzers_names

    def _get_analyzers_mapping(self) -> Dict[str, Analyzer]:
        """
        Get the mapping dict between the name and the instance of the analyzer
        
        Args:
            None

        Returns:
            Dict[str, Analyzer]:A map containing name-analyzer pairs.

        """
        from paddlets import analysis
        analyzers_mapping = {}
        for name, obj in inspect.getmembers(analysis, inspect.isclass):
            if name == self.__class__.__name__:
                continue
            analyzers_mapping[obj.get_properties()['name']] = obj

        return analyzers_mapping

    def _validate_analyzers_names(self, names: Union[str, List[str]] = None) -> None:
        """
        Validate the names of analyzer input by the user 
        If the analyzer names entered by the user do not exist in the library, an error will be reported
        
        Args:
            names(str|List(str)):Names of analyzers

        Returns:
            None

        Raise:
            ValueError

        """
        analyzer_names = self.get_all_analyzers_names(log=False)

        missing_names = set(names) - set(analyzer_names)
        raise_if_not(len(missing_names) == 0,
            f"Invalid analyzer names, analyzer {missing_names} do not exist, please use get_all_analyzers_names() method to get currently supported analyzers!")
